# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.
import datetime

import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
import os
from glob import glob


def read_mfr_txt(projnm):

    # get project information, from txt file, first N lines
    # get MFR data and input into a dictionary
    file = open(projnm, "r")
    items = {}
    count = 0

    for line in file:
        num = line[0]
        print(f' type of first object is {str(type(num))}')
        if num.isnumeric():
            num, loc, desc, photo, action, date = line.strip().split("\t")
            items[num] = (loc, desc, photo, action, date)
        else:
            # do nothing
            count = count + 1
            print(count)

    file.close()
    return items


def get_action_items(lst):
    new_lst = []
    for key in lst.keys():
        if lst[key][3] == str('Yes'):
            print("success")
            new_lst.append(tuple(key) + lst[key])
    return new_lst


def create_doc(projnum):

    script_dir = os.path.dirname(__file__)
    file = open(os.path.join(script_dir, projnum, '{0}_{1}.txt'.format(projnum, 'MFR')))
    itms = {}
    count = 0
    for line in file:
        if count < 4:
            if count == 0:
                #value is project
                x, proj = line.strip().split("\t")
                print(proj)
            if count == 1:
                #value is project number
                y, projnum = line.strip().split("\t")
                print(projnum)
            if count == 2:
                # value is engineer of record
                y, eor = line.strip().split("\t")
                print(eor)
            if count == 3:
                # value is distribution
                # should add functionality for multiple people in distribution list
                y, dist1 = line.strip().split("\t")
                print(dist1)
            count = count + 1
            print(count)
        else:
            # this may be wrong and need to be updated for this function to work
            num, loc, desc, photo, action, date = line.strip().split("\t")
            itms[num] = [loc, desc, photo, action, date]

    num_items = len(itms)
    file.close()

    print(f'Items are {itms}')

    new_list = []
    temp = []

    print(f'itms.items is {itms.items()}')

    for key, val in itms.items():
        print(f' key is {key}')
        print(f' val is {val}')
        val.insert(0,key)
        print(f' val is {val}')
        new_list.append(val)

    print(f'new list is {new_list}')

    document = Document('mfr_temp.docx')

    header = document.sections[0].header
    header.add_paragraph('Project: ' + proj + '\t' + 'Generated: ' + str(datetime.datetime.today()))

    document.add_heading('Mechanical Field Report', 0)

    p1 = document.add_paragraph('Project Name: ', style='Normal')
    p1.add_run(proj)
    p1.add_run('\n')

    p2 = document.add_paragraph('Project Number: ', style='Normal')
    p2.add_run(projnum)
    p2.add_run('\n')

    p3 = document.add_paragraph('Engineer of Record: ', style='Normal')
    p3.add_run(eor)
    p3.add_run(', P.Eng').bold = True
    p3.add_run('\n')

    p4 = document.add_paragraph('Distribution: ', style='Normal')
    p4.add_run(dist1)
    p4.add_run('\n')

    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Location'
    hdr_cells[2].text = 'Description'
    hdr_cells[3].text = 'Photograph'
    hdr_cells[4].text = 'Action'
    for num, loc, desc, photo, action, date in new_list:
        row_cells = table.add_row().cells
        row_cells[0].text = num
        row_cells[1].text = loc
        row_cells[2].text = desc
        pic = row_cells[3].paragraphs[0]
        run = pic.add_run()
        font = run.font
        run.add_picture(os.path.join(projnum, '{0}_{1}'.format('IMG',photo)),width=Inches(2.5))
        run.add_break()  # adds a new line
        run.italic = True
        font.size = Pt(8)
        font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

        run.add_text(photo)  # this adds photo filename following the photo

        row_cells[4].text = action

    document.save('mfr1_test.docx')



# Press the green button in the gutter to run the script.
if __name__ == '__something__':
    filedir = 'C:\\Users\\Darcy.Pederson.RPE\\PycharmProjects\\'
    print(filedir)
    project = "21411-N"
    mfrlst = read_mfr_txt(os.path.join(project, project + '_MFR.txt'))
    print(f' list is {mfrlst}')

    print(f'action items list is {get_action_items(mfrlst)}')

# See PyCharm help at https://www.jetbrains.com/help/pycharm/
